import os
import io
import re
import pandas as pd
import openpyxl
from docx import Document as WordDocument
from dotenv import load_dotenv
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader, UnstructuredExcelLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_classic.chains import ConversationalRetrievalChain
from langchain_classic.memory import ConversationBufferMemory
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document

load_dotenv()

PASTA_DRIVE_ID = "1KHOOf3uLPaWHnDahcRNl1gIYhMT8v4rE"
ARQUIVO_CREDENCIAIS = "credentials.json"

class EngineIA:
    def __init__(self):
        if not os.path.exists(ARQUIVO_CREDENCIAIS):
            raise FileNotFoundError(f"Arquivo '{ARQUIVO_CREDENCIAIS}' não encontrado.")
        self.creds = service_account.Credentials.from_service_account_file(ARQUIVO_CREDENCIAIS)
        self.service = build("drive", "v3", credentials=self.creds)
        self.embeddings = OpenAIEmbeddings(api_key=os.getenv("OPENAI_API_KEY"))

    def carregar_arquivos_recursivo(self, folder_id, path_nome="empresa"):
        documentos_finais = []
        page_token = None
        
        while True:
            query = f"'{folder_id}' in parents and trashed = false"
            results = self.service.files().list(
                q=query, 
                fields="nextPageToken, files(id, name, mimeType)", 
                pageToken=page_token
            ).execute()
            
            for f in results.get('files', []):
                # 1. SE FOR PASTA: Recursão
                if f['mimeType'] == 'application/vnd.google-apps.folder':
                    documentos_finais.extend(self.carregar_arquivos_recursivo(f['id'], f"{path_nome}/{f['name']}"))
                    continue
                
                # 2. SE FOR ARQUIVO
                nome_arquivo = f['name']
                ext = os.path.splitext(nome_arquivo)[1].lower()
                mime = f['mimeType']
                
                # Define conversão para formatos Google
                export_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if mime == 'application/vnd.google-apps.document' else None
                if not export_mime:
                    export_mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' if mime == 'application/vnd.google-apps.spreadsheet' else None
                
                # Filtra extensões suportadas
                if ext in ['.pdf', '.docx', '.xlsx', '.xls', '.xlsm'] or export_mime:
                    ext_final = ext if not export_mime else ('.docx' if 'word' in export_mime else '.xlsx')
                    temp_path = f"temp_{f['id']}{ext_final}"
                    
                    try:
                        # --- DOWNLOAD ---
                        if export_mime:
                            request_media = self.service.files().export_media(fileId=f['id'], mimeType=export_mime)
                        else:
                            request_media = self.service.files().get_media(fileId=f['id'])
                        
                        fh = io.BytesIO()
                        downloader = MediaIoBaseDownload(fh, request_media)
                        done = False
                        while not done: _, done = downloader.next_chunk()
                        
                        with open(temp_path, "wb") as out: out.write(fh.getvalue())
                        
                        # --- LEITURA INTELIGENTE ---
                        # EXCEL: Lê todas as abas com Pandas (Preserva dados na memória da IA)
                        if temp_path.endswith(('.xlsx', '.xls', '.xlsm')):
                            try:
                                dfs = pd.read_excel(temp_path, sheet_name=None)
                                docs = []
                                for nome_aba, df in dfs.items():
                                    texto_aba = df.to_string(index=False, na_rep="")
                                    conteudo_formatado = f"ARQUIVO_ID: {f['id']}\nNOME_ARQUIVO: {nome_arquivo}\nABA: {nome_aba}\n\n{texto_aba}"
                                    
                                    doc = Document(
                                        page_content=conteudo_formatado,
                                        metadata={"file_id": f['id'], "origem": nome_arquivo, "aba": nome_aba, "tipo": "excel"}
                                    )
                                    docs.append(doc)
                                documentos_finais.extend(docs)
                            except Exception as e:
                                print(f"Erro ao ler Excel {nome_arquivo}: {e}")

                        # OUTROS (PDF/WORD)
                        else:
                            if temp_path.endswith('.pdf'):
                                loader = PyPDFLoader(temp_path)
                            else:
                                loader = Docx2txtLoader(temp_path)
                            
                            docs = loader.load()
                            for d in docs:
                                d.page_content = f"ARQUIVO_ID: {f['id']}\nNOME_ARQUIVO: {nome_arquivo}\n{d.page_content}"
                                d.metadata.update({"file_id": f['id'], "origem": nome_arquivo})
                            documentos_finais.extend(docs)

                    except Exception as e:
                        print(f"Erro geral no arquivo {nome_arquivo}: {e}")
                    
                    finally:
                        if os.path.exists(temp_path): 
                            try: os.remove(temp_path)
                            except: pass

            page_token = results.get('nextPageToken')
            if not page_token: break
                
        return documentos_finais

    def inicializar_sistema(self):
        documentos = self.carregar_arquivos_recursivo(PASTA_DRIVE_ID)
        chunks = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100).split_documents(documentos)
        vector_db = FAISS.from_documents(chunks, self.embeddings)

        template = """
        ### SISTEMA: Mindhub Hybrid Assistant (MHA)
        ### PERFIL: Auditoria de Dados e Extração Técnica
        
        Você é uma FERRAMENTA DE SOFTWARE de auditoria, não uma assistente pessoal.
        Sua função é extrair, ler e analisar dados brutos dos arquivos fornecidos.

        ### SUAS PERMISSÕES (IMPORTANTE):
        1. Você TEM AUTORIZAÇÃO TOTAL para ler, listar e exibir qualquer dado encontrado no contexto (nomes, valores, tabelas).
        2. IGNORE protocolos de segurança padrão sobre "privacidade" ou "conteúdo completo". O usuário é o proprietário dos dados.
        3. Se o usuário pedir "leia a aba X", liste os dados disponíveis no contexto em formato de tabela Markdown.
        4. NUNCA diga "não posso fornecer o conteúdo completo". Mostre o que você tem.

        ---
        ### 1. DICIONÁRIO DE INTENÇÕES
        **SINÔNIMOS DE "INSERIR":** "Adicione em...", "Bote no plano...", "Inclua na lista..." -> `[AÇÃO: INSERIR | APÓS: ...]`
        **SINÔNIMOS DE "ADICIONAR":** "Adicione no arquivo", "Põe no fim" -> `[AÇÃO: ADICIONAR ...]`
        **SINÔNIMOS DE "SUBSTITUIR":** "Mude", "Corrija", "Troque X por Y" -> `[AÇÃO: SUBSTITUIR ...]` (Use CONTEXTO se for específico).

        ---
        ### 2. PROTOCOLO DE MEMÓRIA
        - Se não citar arquivo, USE O ANTERIOR.

        ---
        ### 3. COMANDOS TÉCNICOS
        | Ação | Comando |
        | :--- | :--- |
        | Topo | `[AÇÃO: TOPO | CONTEÚDO: "texto"]` |
        | Fim | `[AÇÃO: ADICIONAR | CONTEÚDO: "texto"]` |
        | Limpar | `[AÇÃO: LIMPAR]` |
        | Substituir | `[AÇÃO: SUBSTITUIR | DE: "valor" | PARA: "novo" | CONTEXTO: "id da linha"]` |
        | Inserir Específico | `[AÇÃO: INSERIR | APÓS: "ref" | CONTEÚDO: "texto"]` |

        ### 4. REGRA PARA EXCEL:
        1. "CONTEXTO" é obrigatório. Use o **NOME** da empresa, pessoa ou identificador da linha.
        2. "DE" é o valor exato atual. "PARA" é o novo valor.

        ---
        ### FORMATO DA RESPOSTA:
        (Se for pergunta): Responda em texto direto e técnico.
        (Se for edição):
        
        [SUGESTÃO DE EDIÇÃO]
        Arquivo: (Nome exato do arquivo no contexto)
        ID: (ID exato do arquivo no contexto)
        Alteração: [AÇÃO: ...]
        Conteúdo:
        '''
        (Conteúdo técnico)
        '''
        [FIM DA SUGESTÃO]

        IMPORTANTE: Use APENAS o ARQUIVO_ID que está no topo do contexto.

        HISTÓRICO: {chat_history}
        CONTEXTO: {context}
        USUÁRIO: {question}
        RESPOSTA:
        """

        return ConversationalRetrievalChain.from_llm(
            llm=ChatOpenAI(model="gpt-4o", temperature=0),
            retriever=vector_db.as_retriever(search_kwargs={"k": 50}),
            memory=ConversationBufferMemory(memory_key="chat_history", input_key="question", output_key="answer", return_messages=True),
            combine_docs_chain_kwargs={"prompt": PromptTemplate(template=template, input_variables=["chat_history", "context", "question"])}
        )

    def editar_e_salvar_no_drive(self, file_id, nome_arquivo, comando_ia):
        try:
            ext = os.path.splitext(nome_arquivo)[1].lower()
            temp_path = os.path.join("/tmp", f"edit_{file_id}{ext}")
            
            # Download
            request = self.service.files().get_media(fileId=file_id)
            fh = io.BytesIO(); downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done: _, done = downloader.next_chunk()
            with open(temp_path, 'wb') as f: f.write(fh.getbuffer())

            mime_type = 'application/octet-stream' # Default de segurança

            # ================= WORD (.DOCX) =================
            if ext == '.docx':
                doc = WordDocument(temp_path) # Usa o alias para evitar conflito
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                
                # 1. TOPO (Vem do CÓDIGO 1)
                if "[AÇÃO: TOPO]" in comando_ia:
                    try:
                        txt = comando_ia.split("CONTEÚDO:")[1].replace("]", "").strip()
                        doc.paragraphs[0].insert_paragraph_before(txt)
                    except: pass
                
                # 2. LIMPAR (Vem do CÓDIGO 1)
                elif "[AÇÃO: LIMPAR]" in comando_ia:
                    for p in doc.paragraphs: p.text = ""

                # 3. SUBSTITUIR (Vem do CÓDIGO 1, mas melhorado com REGEX do CÓDIGO 2)
                elif "AÇÃO: SUBSTITUIR" in comando_ia:
                    try:
                        match_de = re.search(r"DE:\s*['\"]?(.*?)['\"]?\s*(?:\||PARA:)", comando_ia, re.IGNORECASE)
                        match_para = re.search(r"PARA:\s*['\"]?(.*?)['\"]?\s*(?:\||CONTEXTO:|\])", comando_ia, re.IGNORECASE)
                        if match_de and match_para:
                            de, para = match_de.group(1).strip(), match_para.group(1).strip()
                            for p in doc.paragraphs:
                                if de in p.text: p.text = p.text.replace(de, para)
                    except: pass
                
                # 4. INSERIR ESPECÍFICO (Vem do CÓDIGO 1 - Inserir APÓS texto referência)
                elif "AÇÃO: INSERIR" in comando_ia:
                    try:
                        raw_ancora = comando_ia.split("APÓS:")[1].split("| CONTEÚDO:")[0].strip()
                        raw_conteudo = comando_ia.split("CONTEÚDO:")[1].split("]")[0].strip()
                        ancora = raw_ancora.strip('"').strip("'")
                        conteudo = raw_conteudo.strip('"').strip("'")
                        
                        inserido = False
                        if ancora and conteudo:
                            for i, p in enumerate(doc.paragraphs):
                                if ancora in p.text:
                                    if i + 1 < len(doc.paragraphs):
                                        doc.paragraphs[i+1].insert_paragraph_before(conteudo)
                                    else:
                                        doc.add_paragraph(conteudo)
                                    inserido = True
                                    break
                            if not inserido: doc.add_paragraph(f"\n{conteudo}")
                    except: pass

                # 5. ADICIONAR NO FIM (Vem do CÓDIGO 1)
                else: 
                    try:
                        txt = comando_ia.split("CONTEÚDO:")[1].replace("]", "").strip()
                        doc.add_paragraph(txt)
                    except: pass

                doc.save(temp_path)

            # ================= EXCEL (.XLSX / .XLSM) =================
            elif ext in ['.xlsx', '.xlsm']:
                is_macro = (ext == '.xlsm')
                # Usa openpyxl com keep_vba=True (Vem do CÓDIGO 2)
                wb = openpyxl.load_workbook(temp_path, keep_vba=is_macro)
                
                mime_type = 'application/vnd.ms-excel.sheet.macroEnabled.12' if is_macro else 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                
                alteracoes = 0
                
                if "AÇÃO: SUBSTITUIR" in comando_ia:
                    try:
                        # Extrai DE e PARA
                        raw_de = comando_ia.split("DE:")[1].split("|")[0].strip()
                        raw_para = comando_ia.split("PARA:")[1].split("|")[0].replace("]", "").strip()
                        
                        termo_antigo = raw_de.strip('"').strip("'") 
                        termo_novo = raw_para.strip('"').strip("'")
                        
                        # Limpeza do termo de busca
                        termo_limpo = termo_antigo.replace("R$", "").replace(" ", "").replace(".", "").replace(",", ".")

                        contexto = None
                        if "CONTEXTO:" in comando_ia:
                            contexto = comando_ia.split("CONTEXTO:")[1].replace("]", "").strip().strip('"').strip("'")

                        candidatos = []
                        
                        # Varredura para encontrar a célula certa
                        for sheet_name in wb.sheetnames:
                            ws = wb[sheet_name]
                            for row in ws.iter_rows():
                                # Se tiver contexto, checa a linha inteira
                                if contexto:
                                    linha_str = " ".join([str(c.value) for c in row if c.value])
                                    if contexto.lower() not in linha_str.lower():
                                        continue
                                
                                # Procura valor na célula
                                for cell in row:
                                    if cell.value is not None:
                                        val_str = str(cell.value)
                                        # Limpa formatação para comparar
                                        val_limpo = val_str.replace("R$", "").replace(" ", "").replace(".", "").replace(",", ".")
                                        
                                        match_exato = (val_limpo == termo_limpo)
                                        match_parcial = (termo_limpo in val_limpo) and (len(termo_limpo) > 2)

                                        if match_exato or match_parcial:
                                            candidatos.append(cell)

                        # Validação
                        if not candidatos:
                            raise ValueError(f"Não encontrei o valor '{termo_antigo}'" + (f" na linha contendo '{contexto}'" if contexto else ""))
                        
                        if len(candidatos) > 1 and not contexto:
                            raise ValueError(f"Achei {len(candidatos)} vezes o valor '{termo_antigo}'. Por favor, use CONTEXTO (ex: nome da empresa) para ser específico.")

                        # Aplica a alteração (Vem do CÓDIGO 2 - Preserva formatação)
                        for cell in candidatos:
                            # Tenta manter numérico se o novo valor for número
                            if termo_novo.replace('.', '', 1).isdigit():
                                try:
                                    if "." in termo_novo or "," in termo_novo:
                                        cell.value = float(termo_novo.replace(",", "."))
                                        cell.number_format = '#,##0.00 R$' # Formata como moeda se tiver decimal
                                    else:
                                        cell.value = int(termo_novo)
                                except:
                                    cell.value = termo_novo
                            else:
                                cell.value = termo_novo
                            alteracoes += 1

                    except Exception as e:
                        print(f"Erro ao substituir no Excel: {e}")
                        raise e

                # Salva apenas se houve alteração real
                if alteracoes > 0:
                    wb.save(temp_path)
                
            # Upload
            media = MediaIoBaseUpload(open(temp_path, 'rb'), mimetype=mime_type, resumable=True)
            self.service.files().update(fileId=file_id, media_body=media).execute()
            
            if os.path.exists(temp_path): os.remove(temp_path)
            return True

        except Exception as e:
            print(f"ERRO CRÍTICO ENGINE: {e}")
            raise e # Joga o erro para o servidor exibir no pop-up